"""Build the official graduation report DOCX from its Markdown source."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "bao_cao_do_an_tot_nghiep.md"
OUTPUT = ROOT / "docs" / "bao_cao_do_an_tot_nghiep.docx"

BODY_FONT = "Times New Roman"
CODE_FONT = "Consolas"
BLACK = RGBColor(0, 0, 0)
MUTED = RGBColor(85, 85, 85)
LIGHT_GRAY = "F2F2F2"
VERY_LIGHT_GRAY = "F8F8F8"
CONTENT_WIDTH_DXA = 8791
TABLE_INDENT_DXA = 120

FIGURE_NOTE_RE = re.compile(
    r"^>\s*\*\*Gợi ý chèn hình\s+([0-9.]+)\s*-\s*(.+?):\*\*\s*(.+)$",
    re.IGNORECASE,
)
TABLE_CAPTION_RE = re.compile(
    r"^\*\*Bảng\s+([0-9A-Z]+\.[0-9]+)\.\s+(.+)\*\*$",
    re.IGNORECASE,
)
INLINE_TOKEN_RE = re.compile(r"(`[^`\n]+`|\*\*.+?\*\*|https?://[^\s]+|\[(\d+)\])")


def set_run_font(
    run,
    *,
    name: str = BODY_FONT,
    size: float = 12,
    bold: bool | None = None,
    italic: bool | None = None,
    color: RGBColor = BLACK,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_paragraph_border(paragraph, *, color: str = "B7B7B7", size: int = 6, space: int = 4) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), str(space))
        tag.set(qn("w:color"), color)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table_widths(column_count: int) -> list[int]:
    presets = {
        1: [1.0],
        2: [0.31, 0.69],
        3: [0.22, 0.34, 0.44],
        4: [0.18, 0.24, 0.24, 0.34],
        5: [0.15, 0.18, 0.20, 0.20, 0.27],
    }
    ratios = presets.get(column_count, [1 / column_count] * column_count)
    widths = [round(CONTENT_WIDTH_DXA * ratio) for ratio in ratios]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_hyperlink(paragraph, text: str, url: str):
    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run_element = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), BODY_FONT)
    fonts.set(qn("w:hAnsi"), BODY_FONT)
    fonts.set(qn("w:eastAsia"), BODY_FONT)
    r_pr.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "24")
    r_pr.append(size)
    run_element.append(r_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run_element.append(text_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)


def add_inline(paragraph, text: str, *, size: float = 12, color: RGBColor = BLACK) -> None:
    position = 0
    for match in INLINE_TOKEN_RE.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name=CODE_FONT, size=max(9.5, size - 1), color=color)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True, color=color)
        elif token.startswith("http"):
            trailing = ""
            while token and token[-1] in ".,;:":
                trailing = token[-1] + trailing
                token = token[:-1]
            add_hyperlink(paragraph, token, token)
            if trailing:
                run = paragraph.add_run(trailing)
                set_run_font(run, size=size, color=color)
        else:
            run = paragraph.add_run(token)
            set_run_font(run, size=size, color=color)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size=size, color=color)


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction_text, separate, display, end])


def set_page_number_format(section, *, fmt: str, start: int = 1) -> None:
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:pgNumType"))
    if existing is not None:
        sect_pr.remove(existing)
    page_num_type = OxmlElement("w:pgNumType")
    page_num_type.set(qn("w:fmt"), fmt)
    page_num_type.set(qn("w:start"), str(start))
    sect_pr.append(page_num_type)


def set_section_geometry(section) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)


def clear_container(container) -> None:
    for paragraph in container.paragraphs:
        for run in paragraph.runs:
            run._element.getparent().remove(run._element)
    for table in list(container.tables):
        table._element.getparent().remove(table._element)


def set_front_footer(section) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    clear_container(section.header)
    clear_container(section.footer)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(footer, "PAGE", "i")
    for run in footer.runs:
        set_run_font(run, size=10)


def set_main_header_footer(section) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    clear_container(section.header)
    clear_container(section.footer)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_field(header, 'STYLEREF "Heading 1"', "BÁO CÁO ĐỒ ÁN TỐT NGHIỆP")
    for run in header.runs:
        set_run_font(run, size=10)
    p_pr = header._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "808080")
    borders.append(bottom)
    p_pr.append(borders)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(footer, "PAGE", "1")
    for run in footer.runs:
        set_run_font(run, size=10)


def set_blank_section(section) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    clear_container(section.header)
    clear_container(section.footer)


def set_style_language(style, language: str = "vi-VN") -> None:
    r_pr = style._element.get_or_add_rPr()
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), language)
    lang.set(qn("w:eastAsia"), language)


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(12)
    normal.font.color.rgb = BLACK
    set_style_language(normal)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Cm(1.27)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    heading_tokens = {
        "Heading 1": (16, 0, 18, WD_ALIGN_PARAGRAPH.CENTER),
        "Heading 2": (14, 14, 6, WD_ALIGN_PARAGRAPH.LEFT),
        "Heading 3": (13, 10, 5, WD_ALIGN_PARAGRAPH.LEFT),
        "Heading 4": (12, 8, 4, WD_ALIGN_PARAGRAPH.LEFT),
    }
    for style_name, (size, before, after, alignment) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLACK
        set_style_language(style)
        style.paragraph_format.alignment = alignment
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.line_spacing = 1.15

    heading1 = styles["Heading 1"]
    heading1.paragraph_format.page_break_before = True

    caption = styles["Caption"]
    caption.font.name = BODY_FONT
    caption._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    caption.font.size = Pt(11)
    caption.font.italic = False
    caption.font.color.rgb = BLACK
    set_style_language(caption)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = False

    if "Table Caption" not in styles:
        table_caption = styles.add_style("Table Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        table_caption = styles["Table Caption"]
    table_caption.font.name = BODY_FONT
    table_caption._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    table_caption._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    table_caption._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    table_caption.font.size = Pt(11)
    table_caption.font.bold = True
    table_caption.font.color.rgb = BLACK
    set_style_language(table_caption)
    table_caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_caption.paragraph_format.first_line_indent = Cm(0)
    table_caption.paragraph_format.space_before = Pt(8)
    table_caption.paragraph_format.space_after = Pt(4)
    table_caption.paragraph_format.keep_with_next = True
    table_caption.paragraph_format.keep_together = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(12)
        set_style_language(style)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style.paragraph_format.left_indent = Cm(0.95)
        style.paragraph_format.first_line_indent = Cm(-0.49)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    if "Code Block" not in styles:
        code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = styles["Code Block"]
    code_style.font.name = CODE_FONT
    code_style._element.rPr.rFonts.set(qn("w:ascii"), CODE_FONT)
    code_style._element.rPr.rFonts.set(qn("w:hAnsi"), CODE_FONT)
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), CODE_FONT)
    code_style.font.size = Pt(9.5)
    code_style.paragraph_format.left_indent = Cm(0.25)
    code_style.paragraph_format.right_indent = Cm(0.25)
    code_style.paragraph_format.first_line_indent = Cm(0)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(6)
    code_style.paragraph_format.line_spacing = 1.0

    if "Figure Placeholder" not in styles:
        figure_style = styles.add_style("Figure Placeholder", WD_STYLE_TYPE.PARAGRAPH)
    else:
        figure_style = styles["Figure Placeholder"]
    figure_style.font.name = BODY_FONT
    figure_style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    figure_style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    figure_style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    figure_style.font.size = Pt(10.5)
    figure_style.font.color.rgb = BLACK
    set_style_language(figure_style)
    figure_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure_style.paragraph_format.first_line_indent = Cm(0)
    figure_style.paragraph_format.left_indent = Cm(0.5)
    figure_style.paragraph_format.right_indent = Cm(0.5)
    figure_style.paragraph_format.space_before = Pt(8)
    figure_style.paragraph_format.space_after = Pt(2)
    figure_style.paragraph_format.line_spacing = 1.15
    figure_style.paragraph_format.keep_together = True
    figure_style.paragraph_format.keep_with_next = True


def add_cover(document: Document) -> None:
    section = document.sections[0]
    set_section_geometry(section)
    set_blank_section(section)

    def centered(text: str, *, size: float, bold: bool = False, before: float = 0, after: float = 0):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.space_before = Pt(before)
        paragraph.paragraph_format.space_after = Pt(after)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(text)
        set_run_font(run, size=size, bold=bold)
        return paragraph

    centered("[TÊN TRƯỜNG]", size=13, bold=True, after=10)
    centered("[TÊN KHOA/BỘ MÔN]", size=13, bold=True)
    centered("BÁO CÁO ĐỒ ÁN TỐT NGHIỆP", size=18, bold=True, before=72, after=18)
    centered("CHUYÊN NGÀNH: [TÊN CHUYÊN NGÀNH]", size=14, bold=True)
    centered("TÊN ĐỀ TÀI", size=14, bold=True, before=50, after=12)
    centered(
        "ỨNG DỤNG AI HỖ TRỢ PHÁT HIỆN VÀ GỢI Ý\n"
        "LÀM SẠCH DỮ LIỆU THIẾU, SAI LUẬT VÀ BẤT\n"
        "THƯỜNG TRONG CƠ SỞ DỮ LIỆU GIAO DỊCH\n"
        "THƯƠNG MẠI ĐIỆN TỬ",
        size=16,
        bold=True,
    )

    metadata = [
        ("Sinh viên thực hiện:", "[HỌ VÀ TÊN SINH VIÊN]"),
        ("Mã số sinh viên:", "[MÃ SỐ SINH VIÊN]"),
        ("Giảng viên hướng dẫn:", "[TÊN GIẢNG VIÊN HƯỚNG DẪN]"),
    ]
    first = True
    for label, value in metadata:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Cm(4.0)
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.space_before = Pt(80 if first else 2)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.0
        label_run = paragraph.add_run(f"{label:<25}")
        set_run_font(label_run, size=12, bold=True)
        value_run = paragraph.add_run(value)
        set_run_font(value_run, size=12)
        first = False

    centered("[ĐỊA ĐIỂM], [THÁNG, NĂM]", size=12, bold=True, before=90, after=8)
    centered("Năm học: [NĂM HỌC]", size=12)


def add_two_blank_pages_after_cover(document: Document) -> None:
    document.add_page_break()
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.add_run("")
    document.add_page_break()
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.add_run("")


def start_front_matter(document: Document):
    section = document.add_section(WD_SECTION.NEW_PAGE)
    set_section_geometry(section)
    set_page_number_format(section, fmt="lowerRoman", start=1)
    set_front_footer(section)
    return section


def start_main_matter(document: Document):
    section = document.add_section(WD_SECTION.NEW_PAGE)
    set_section_geometry(section)
    set_page_number_format(section, fmt="decimal", start=1)
    set_main_header_footer(section)
    return section


def add_two_blank_pages_at_end(document: Document) -> None:
    section = document.add_section(WD_SECTION.NEW_PAGE)
    set_section_geometry(section)
    set_blank_section(section)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.add_run("")
    document.add_page_break()
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.add_run("")


def add_navigation_fields(document: Document) -> None:
    heading = document.add_paragraph("MỤC LỤC", style="Heading 1")
    heading.paragraph_format.page_break_before = True
    toc = document.add_paragraph()
    toc.paragraph_format.first_line_indent = Cm(0)
    toc.paragraph_format.line_spacing = 1.0
    add_field(toc, 'TOC \\o "1-3" \\h \\z \\u', "Nhấp chuột phải và chọn Update Field để cập nhật mục lục.")

    figure_heading = document.add_paragraph("DANH MỤC HÌNH", style="Heading 1")
    figure_heading.paragraph_format.page_break_before = True
    figure_toc = document.add_paragraph()
    figure_toc.paragraph_format.first_line_indent = Cm(0)
    figure_toc.paragraph_format.line_spacing = 1.0
    add_field(
        figure_toc,
        'TOC \\h \\z \\t "Caption,1"',
        "Nhấp chuột phải và chọn Update Field để cập nhật danh mục hình.",
    )

    table_heading = document.add_paragraph("DANH MỤC BẢNG", style="Heading 1")
    table_heading.paragraph_format.page_break_before = True
    table_toc = document.add_paragraph()
    table_toc.paragraph_format.first_line_indent = Cm(0)
    table_toc.paragraph_format.line_spacing = 1.0
    add_field(
        table_toc,
        'TOC \\h \\z \\t "Table Caption,1"',
        "Nhấp chuột phải và chọn Update Field để cập nhật danh mục bảng.",
    )


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_separator_row(cells: list[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def add_table(document: Document, lines: list[str]) -> None:
    rows = [parse_table_row(line) for line in lines]
    if len(rows) > 1 and is_separator_row(rows[1]):
        rows.pop(1)
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    normalized = [row + [""] * (column_count - len(row)) for row in rows]
    table = document.add_table(rows=len(normalized), cols=column_count)
    table.style = "Table Grid"
    widths = table_widths(column_count)
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])

    for row_index, row in enumerate(normalized):
        for column_index, text in enumerate(row):
            cell = table.cell(row_index, column_index)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline(paragraph, text, size=10.5)
            if row_index == 0:
                set_cell_shading(cell, LIGHT_GRAY)
                for run in paragraph.runs:
                    run.bold = True
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    spacer.paragraph_format.first_line_indent = Cm(0)


def add_code_block(document: Document, language: str, code_lines: list[str]) -> None:
    paragraph = document.add_paragraph(style="Code Block")
    paragraph.paragraph_format.keep_together = len(code_lines) <= 12
    label = paragraph.add_run(language or "text")
    set_run_font(label, name=CODE_FONT, size=9, bold=True, color=MUTED)
    label.add_break()
    for index, line in enumerate(code_lines):
        run = paragraph.add_run(line)
        set_run_font(run, name=CODE_FONT, size=9.5)
        if index < len(code_lines) - 1:
            run.add_break()
    set_paragraph_shading(paragraph, VERY_LIGHT_GRAY)
    set_paragraph_border(paragraph, color="B7B7B7", size=4, space=4)


def add_table_caption(document: Document, number: str, title: str) -> None:
    caption = document.add_paragraph(style="Table Caption")
    run = caption.add_run(f"Bảng {number}. {title}")
    set_run_font(run, size=11, bold=True)


def add_figure_placeholder(document: Document, number: str, title: str, description: str) -> None:
    paragraph = document.add_paragraph(style="Figure Placeholder")
    heading = paragraph.add_run(f"VỊ TRÍ CHỜ HÌNH {number}\n")
    set_run_font(heading, size=11, bold=True)
    add_inline(paragraph, description, size=10.5)
    set_paragraph_shading(paragraph, "FAFAFA")
    set_paragraph_border(paragraph, color="808080", size=6, space=8)

    caption = document.add_paragraph(style="Caption")
    run = caption.add_run(f"Hình {number}. {title}")
    set_run_font(run, size=11)


def add_heading(document: Document, level: int, title: str, *, first_on_page: bool = False) -> None:
    style_level = min(level, 4)
    paragraph = document.add_paragraph(title, style=f"Heading {style_level}")
    if first_on_page:
        paragraph.paragraph_format.page_break_before = False
    for run in paragraph.runs:
        set_run_font(run, size={1: 16, 2: 14, 3: 13, 4: 12}[style_level], bold=True)


def add_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_inline(paragraph, text)


def add_list_item(document: Document, text: str, *, ordered: bool) -> None:
    paragraph = document.add_paragraph(style="List Number" if ordered else "List Bullet")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_inline(paragraph, text)


def add_block_quote(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(1)
    paragraph.paragraph_format.right_indent = Cm(0.5)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_inline(paragraph, text, size=11, color=MUTED)
    for run in paragraph.runs:
        run.italic = True


def convert_markdown(document: Document, source: str) -> None:
    lines = source.splitlines()
    index = 0
    started = False
    navigation_added = False
    main_started = False
    appendix_mode = False
    first_front_heading = True
    first_main_heading = True

    while index < len(lines):
        line = lines[index].rstrip()

        if not started:
            if line.strip() == "# Lời cam đoan":
                started = True
            else:
                index += 1
                continue

        if line.startswith("```"):
            language = line[3:].strip()
            code_lines: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].startswith("```"):
                code_lines.append(lines[index].rstrip())
                index += 1
            add_code_block(document, language, code_lines)
            index += 1
            continue

        if line.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].lstrip().startswith("|"):
                table_lines.append(lines[index].rstrip())
                index += 1
            add_table(document, table_lines)
            continue

        table_caption_match = TABLE_CAPTION_RE.match(line)
        if table_caption_match:
            add_table_caption(document, *table_caption_match.groups())
            index += 1
            continue

        figure_match = FIGURE_NOTE_RE.match(line)
        if figure_match:
            add_figure_placeholder(document, *figure_match.groups())
            index += 1
            continue

        if line.startswith("# "):
            title = line[2:].strip()
            if title == "PHỤ LỤC":
                appendix_mode = True
                index += 1
                continue

            if title.startswith("CHƯƠNG 1.") and not main_started:
                start_main_matter(document)
                main_started = True
                first_main_heading = True

            if title == "Danh mục từ viết tắt" and not navigation_added:
                add_navigation_fields(document)
                navigation_added = True

            first_on_page = False
            if not main_started and first_front_heading:
                first_on_page = True
                first_front_heading = False
            elif main_started and first_main_heading:
                first_on_page = True
                first_main_heading = False
            add_heading(document, 1, title, first_on_page=first_on_page)
            index += 1
            continue

        if line.startswith("## "):
            title = line[3:].strip()
            if appendix_mode and title.lower().startswith("phụ lục"):
                add_heading(document, 1, title)
            else:
                add_heading(document, 2, title)
            index += 1
            continue

        if line.startswith("### "):
            add_heading(document, 3, line[4:].strip())
            index += 1
            continue

        if line.startswith("#### "):
            add_heading(document, 4, line[5:].strip())
            index += 1
            continue

        if not line or line == "---":
            index += 1
            continue

        ordered = re.match(r"^\d+\.\s+(.+)$", line)
        if ordered:
            add_list_item(document, ordered.group(1), ordered=True)
            index += 1
            continue

        unordered = re.match(r"^-\s+(.+)$", line)
        if unordered:
            add_list_item(document, unordered.group(1), ordered=False)
            index += 1
            continue

        if line.startswith("> "):
            add_block_quote(document, line[2:].strip())
            index += 1
            continue

        add_body_paragraph(document, line)
        index += 1


def request_field_updates(document: Document) -> None:
    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def set_document_properties(document: Document) -> None:
    properties = document.core_properties
    properties.title = (
        "Ứng dụng AI hỗ trợ phát hiện và gợi ý làm sạch dữ liệu thiếu, "
        "sai luật và bất thường trong cơ sở dữ liệu giao dịch thương mại điện tử"
    )
    properties.subject = "Báo cáo đồ án tốt nghiệp"
    properties.author = "[HỌ VÀ TÊN SINH VIÊN]"
    properties.keywords = "data cleaning, anomaly detection, Isolation Forest, PostgreSQL"


def main() -> None:
    source = SOURCE.read_text(encoding="utf-8")
    document = Document()
    configure_styles(document)
    set_document_properties(document)
    add_cover(document)
    add_two_blank_pages_after_cover(document)
    start_front_matter(document)
    convert_markdown(document, source)
    add_two_blank_pages_at_end(document)
    request_field_updates(document)
    document.save(OUTPUT)
    print(f"Generated {OUTPUT}")


if __name__ == "__main__":
    main()
